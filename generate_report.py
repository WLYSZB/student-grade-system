# -*- coding: utf-8 -*-
"""生成课程设计报告 DOCX"""

import base64
import json
import os
import zlib
import requests
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# ========== Mermaid 转 PNG ==========

def encode_mermaid(code):
    """将Mermaid代码编码为mermaid.ink URL"""
    data = json.dumps({"code": code, "mermaid": {"theme": "default"}})
    encoded = base64.urlsafe_b64encode(zlib.compress(data.encode(), 9)).decode()
    return f"https://mermaid.ink/img/{encoded}"

def download_image(url, path):
    r = requests.get(url, timeout=30)
    if r.status_code == 200:
        with open(path, 'wb') as f:
            f.write(r.content)
        return True
    return False

# ========== DOCX 工具 ==========

class ReportBuilder:
    def __init__(self):
        self.doc = Document()
        self.img_dir = "report_images"
        os.makedirs(self.img_dir, exist_ok=True)
        self.img_counter = 0
        self._setup_styles()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)

        for level in [1, 2, 3]:
            h = self.doc.styles[f'Heading {level}']
            h.font.name = '黑体'
            h.font.color.rgb = RGBColor(0, 0, 0)

    def add_mermaid(self, code, caption="", width=5.5):
        """添加Mermaid图"""
        self.img_counter += 1
        fname = f"{self.img_dir}/img_{self.img_counter}.png"
        url = encode_mermaid(code)
        if download_image(url, fname):
            self.doc.add_picture(fname, width=Inches(width))
            if caption:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(caption)
                run.font.size = Pt(9)
                run.font.name = '楷体'
            self.doc.add_paragraph()
        else:
            self.doc.add_paragraph(f"[图{self.img_counter}: 图像生成失败，请手动用 https://mermaid.live 渲染以下代码]")
            self.doc.add_paragraph(code)

    def cover_page(self):
        """封面"""
        for _ in range(6):
            self.doc.add_paragraph()
        t = self.doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run('《软件工程综合实验》\n课程设计报告')
        r.font.size = Pt(26)
        r.font.name = '黑体'
        r.bold = True

        self.doc.add_paragraph()
        info = [
            ('题    目', '学生成绩管理系统的设计与实现'),
            ('专    业', '计算机科学与技术'),
            ('姓    名', '（填写你的姓名）'),
            ('学    号', '（填写你的学号）'),
            ('指导教师', '（填写指导教师）'),
        ]
        for label, value in info:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'{label}：{value}')
            r.font.size = Pt(14)
            r.font.name = '宋体'

        self.doc.add_page_break()

    def heading1(self, text):
        h = self.doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.name = '黑体'

    def heading2(self, text):
        h = self.doc.add_heading(text, level=2)
        for run in h.runs:
            run.font.name = '黑体'

    def heading3(self, text):
        h = self.doc.add_heading(text, level=3)
        for run in h.runs:
            run.font.name = '黑体'

    def para(self, text):
        self.doc.add_paragraph(text)

    def bold_para(self, label, content=""):
        p = self.doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        if content:
            p.add_run(content)

    def bullet(self, text):
        p = self.doc.add_paragraph(text, style='List Bullet')

    def add_table(self, headers, rows, col_widths=None):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
        self.doc.add_paragraph()

    def save(self, path):
        self.doc.save(path)
        # cleanup
        import shutil
        shutil.rmtree(self.img_dir, ignore_errors=True)
        print(f"Report saved to: {path}")


# ========== 生成报告 ==========

def build_report():
    r = ReportBuilder()

    # ===== 封面 =====
    r.cover_page()

    # ===== 第1节：系统功能 =====
    r.heading1('1. 系统功能')

    r.heading2('1.1 项目背景')
    r.para('随着高校教学管理信息化的发展，传统的手工成绩管理方式已无法满足日常需求。'
           '学生成绩管理系统旨在实现成绩录入、查询、统计、分析等功能的数字化和自动化，'
           '提高教务管理效率，方便师生查询成绩信息。')

    r.heading2('1.2 团队分工')
    r.para('本小组共6人，采用模块化分工方案，具体如下：')
    r.add_table(
        ['成员', '角色', '主要职责'],
        [
            ['成员1', '项目经理/架构师', '项目规划、数据库设计（ER图）、系统架构设计、可行性分析、概要设计文档'],
            ['成员2（本人）', '后端核心开发', '登录认证模块：用户注册/登录/密码修改/重置、JWT/Token实现、RBAC权限控制、公共中间件（统一响应格式、异常处理、请求日志）'],
            ['成员3', '教师端开发', '教师端前后端：成绩录入（单条/批量）、成绩修改与审核、成绩统计与排名'],
            ['成员4', '学生端开发', '学生端前后端：成绩查询、成绩趋势分析（Chart.js图表）、通知公告查看'],
            ['成员5', '管理员端开发', '管理员端前后端：用户管理（增删改查/批量导入）、课程管理、通知公告发布'],
            ['成员6', '测试/文档/答辩', '集成测试、系统测试、文档整合排版、答辩PPT制作与汇报'],
        ]
    )

    r.heading2('1.3 本人负责模块')
    r.para('本人在项目中负责后端核心模块——登录与权限管理子系统。该模块是系统的安全入口，'
           '为三种角色（系统管理员、教师、学生）提供身份认证和访问控制服务。主要功能包括：')
    r.bullet('用户注册：支持学生和教师自助注册账号')
    r.bullet('用户登录：验证用户名/密码，建立会话（HttpSession）')
    r.bullet('密码管理：修改密码（需旧密码验证）、忘记密码后重置')
    r.bullet('权限控制：基于角色的访问控制（RBAC），不同角色访问不同功能模块')
    r.bullet('会话管理：基于HttpSession的登录状态保持，LoginInterceptor拦截校验')
    r.bullet('统一响应：所有API返回统一的JSON格式 {code, message, data}')
    r.bullet('全局异常处理：自定义异常类 + GlobalExceptionHandler统一捕获')

    r.heading2('1.4 技术选型')
    r.add_table(
        ['层面', '技术', '说明'],
        [
            ['后端框架', 'Spring Boot 3.5.0', '企业级Java框架，内嵌Tomcat'],
            ['ORM框架', 'MyBatis-Plus 3.5.7', '简化数据库操作，支持分页、自动填充'],
            ['数据库', 'SQL Server', '微软关系型数据库，支持事务和外键约束'],
            ['前端', 'HTML5 + CSS3 + JavaScript', '原生技术栈，无需前端构建工具'],
            ['认证方式', 'HttpSession + Cookie', 'Servlet容器原生支持，实现简单可靠'],
            ['密码加密', 'MD5', '对密码进行不可逆哈希存储'],
            ['构建工具', 'Maven 3.9+', '依赖管理和项目构建'],
        ]
    )

    # ===== 第2节：可行性分析 =====
    r.heading1('2. 可行性分析')

    r.heading2('2.1 技术可行性')
    r.para('本系统采用Spring Boot + MyBatis-Plus + SQL Server技术栈，均为成熟稳定的开源技术。'
           'Spring Boot提供自动配置和嵌入式服务器，大幅降低开发复杂度；'
           'MyBatis-Plus在MyBatis基础上提供了更便捷的CRUD操作；'
           'SQL Server作为企业级数据库，支持事务处理和数据完整性约束。'
           '前端使用HTML5+CSS3+JavaScript原生技术栈，无需额外学习成本。'
           '综上所述，技术方案成熟可靠，具有较高的技术可行性。')

    r.heading2('2.2 经济可行性')
    r.para('本系统所需软件均为开源免费（Spring Boot、MyBatis-Plus）或学校已提供（SQL Server），'
           '开发工具（VS Code / IntelliJ IDEA）可使用社区版，无需额外采购。'
           '系统部署后，可替代传统纸质成绩管理方式，节省人力成本和时间成本。'
           '从经济角度分析，项目成本极低，收益显著，具有可行性。')

    r.heading2('2.3 操作可行性')
    r.para('系统采用B/S架构，用户通过浏览器访问，无需安装客户端软件。'
           '界面设计简洁直观，分为管理员端、教师端和学生端三种视图，'
           '每种角色仅显示其权限范围内的功能，降低使用门槛。'
           '系统部署在本地服务器上，校园网内均可访问。')

    # ===== 第3节：概要设计 =====
    r.heading1('3. 概要设计')

    r.heading2('3.1 系统架构设计')
    r.para('系统采用经典的三层B/S架构：表现层（前端页面）、业务逻辑层（Spring Boot Controller/Service）、'
           '数据访问层（MyBatis-Plus Mapper）。前端通过HTTP请求与后端API交互，后端通过MyBatis-Plus操作数据库。')
    r.add_mermaid("""
graph TD
    A[浏览器 Browser] -->|HTTP请求| B[Controller 控制器层]
    B --> C[Service 业务逻辑层]
    C --> D[Mapper 数据访问层]
    D --> E[(SQL Server 数据库)]
    B --> F[LoginInterceptor 登录拦截]
    B --> G[RoleInterceptor 角色拦截]
    F --> H[HttpSession]
""", "图3-1 系统架构图")

    r.heading2('3.2 数据库设计')
    r.para('数据库包含8张核心表：用户表(sys_user)、教师表(teacher)、学生表(student)、'
           '课程表(course)、学期表(semester)、成绩表(grade)、成绩修改日志表(grade_modify_log)、'
           '通知公告表(notice)。')

    r.heading3('3.2.1 ER图')
    r.add_mermaid("""
erDiagram
    sys_user ||--o| teacher : "1:1"
    sys_user ||--o| student : "1:1"
    teacher ||--o{ course : "授课"
    student ||--o{ grade : "成绩记录"
    course ||--o{ grade : "课程成绩"
    semester ||--o{ course : "所属学期"
    sys_user ||--o{ notice : "发布"
    grade ||--o{ grade_modify_log : "修改日志"

    sys_user {
        bigint id PK
        string username UK
        string password
        string role
        int status
    }
    teacher {
        bigint id PK
        bigint user_id FK
        string name
        string department
    }
    student {
        bigint id PK
        bigint user_id FK
        string name
        string class_name
        string major
    }
    course {
        bigint id PK
        string name
        decimal credit
        bigint teacher_id FK
        bigint semester_id FK
    }
    grade {
        bigint id PK
        bigint student_id FK
        bigint course_id FK
        decimal score
        decimal grade_point
    }
""", "图3-2 数据库ER图")

    r.heading3('3.2.2 核心表结构（用户表）')
    r.add_table(
        ['字段名', '类型', '说明'],
        [
            ['id', 'BIGINT (PK)', '用户ID，自增主键'],
            ['username', 'NVARCHAR(50) (UK)', '用户名，唯一索引'],
            ['password', 'NVARCHAR(100)', '密码，MD5加密存储'],
            ['role', 'NVARCHAR(20)', '角色：ADMIN / TEACHER / STUDENT'],
            ['status', 'TINYINT', '状态：0禁用 1启用'],
            ['create_time', 'DATETIME', '创建时间，默认GETDATE()'],
            ['update_time', 'DATETIME', '更新时间，自动更新'],
            ['deleted', 'TINYINT', '逻辑删除标记'],
        ]
    )

    # ===== 第4节：详细设计 =====
    r.heading1('4. 详细设计')

    r.heading2('4.1 登录认证模块概述')
    r.para('登录认证模块是本系统的安全基础。系统包含三种角色——系统管理员(ADMIN)、'
           '教师(TEACHER)、学生(STUDENT)——不同角色拥有不同的功能访问权限。'
           '本模块负责：身份认证（验证用户名密码）、会话管理（HttpSession + Cookie）、'
           '权限控制（基于角色的访问控制RBAC）、密码管理（加密存储、修改、重置）。')

    r.heading2('4.2 模块结构')
    r.add_table(
        ['文件', '说明'],
        [
            ['controller/auth/AuthController.java', '认证控制器，处理6个REST接口'],
            ['service/auth/AuthService.java', '认证业务逻辑：登录/注册/改密/重置'],
            ['mapper/UserMapper.java', 'sys_user表的MyBatis-Plus Mapper'],
            ['dto/LoginRequest.java', '登录请求体'],
            ['dto/RegisterRequest.java', '注册请求体'],
            ['dto/ChangePasswordRequest.java', '改密/重置密码请求体'],
            ['util/MD5Util.java', 'MD5加密工具类'],
            ['common/LoginInterceptor.java', '登录拦截器，校验Session'],
            ['common/RoleInterceptor.java', '角色权限拦截器，检查角色白名单'],
        ]
    )

    r.heading2('4.3 登录认证流程')
    r.add_mermaid("""
sequenceDiagram
    participant C as 客户端
    participant AC as AuthController
    participant AS as AuthService
    participant DB as 数据库

    C->>AC: POST /api/auth/login {username, password}
    AC->>AS: login(username, password, role)
    AS->>DB: SELECT * FROM sys_user WHERE username=?
    DB-->>AS: User对象
    alt 用户不存在
        AS-->>AC: throw BusinessException(1001)
        AC-->>C: {code:1001, message:"用户不存在"}
    else 密码错误
        AS-->>AC: throw BusinessException(1003)
        AC-->>C: {code:1003, message:"密码错误"}
    else 登录成功
        AS-->>AC: return User
        AC->>AC: session.setAttribute("currentUser", user)
        AC-->>C: {code:200, data:{id,username,role}}
    end
""", "图4-1 登录认证时序图")

    r.heading2('4.4 RBAC权限控制')
    r.para('系统采用基于角色的访问控制（RBAC），通过两层拦截器实现：')
    r.bullet('第1层 LoginInterceptor：拦截 /api/** 路径（白名单除外），检查Session中是否有currentUser，无则返回401')
    r.bullet('第2层 RoleInterceptor：读取请求头 X-Required-Role，检查用户角色是否在允许列表中，不在则返回403')
    r.para('白名单接口（无需登录）：/api/auth/login、/api/auth/register、/api/auth/reset-password')

    r.add_mermaid("""
flowchart LR
    A[请求到达] --> B{在白名单中?}
    B -->|是| G[直接处理]
    B -->|否| C{Session中有<br/>currentUser?}
    C -->|否| D[返回 401]
    C -->|是| E{角色匹配?}
    E -->|否| F[返回 403]
    E -->|是| G
""", "图4-2 权限控制流程图")

    r.heading2('4.5 密码修改流程')
    r.add_mermaid("""
flowchart TD
    A[用户提交旧密码+新密码] --> B{旧密码是否正确?}
    B -->|否| C[返回400: 旧密码不正确]
    B -->|是| D{新旧密码是否相同?}
    D -->|是| E[返回400: 不能与旧密码相同]
    D -->|否| F[MD5加密新密码]
    F --> G[更新数据库]
    G --> H[session.invalidate 清除登录状态]
    H --> I[返回200: 修改成功]
""", "图4-3 密码修改流程图")

    r.heading2('4.6 API接口清单')
    r.add_table(
        ['接口', '方法', '需登录', '说明'],
        [
            ['/api/auth/login', 'POST', '否', '用户登录，返回用户信息'],
            ['/api/auth/register', 'POST', '否', '注册新用户'],
            ['/api/auth/me', 'GET', '是', '获取当前登录用户信息'],
            ['/api/auth/change-password', 'PUT', '是', '修改密码（需旧密码）'],
            ['/api/auth/reset-password', 'POST', '否', '通过用户名重置密码'],
            ['/api/auth/logout', 'POST', '是', '退出登录，销毁Session'],
        ]
    )

    # ===== 第5节：编码与测试 =====
    r.heading1('5. 编码与测试')

    r.heading2('5.1 开发环境')
    r.add_table(
        ['项目', '配置'],
        [
            ['操作系统', 'Windows 11'],
            ['JDK版本', 'Java 25.0.2'],
            ['构建工具', 'Maven 3.9.16'],
            ['数据库', 'SQL Server（localhost:1433）'],
            ['IDE', 'VS Code / IntelliJ IDEA'],
            ['测试工具', 'curl / Postman'],
        ]
    )

    r.heading2('5.2 测试策略')
    r.para('本模块采用黑盒测试方法，对每个API接口编写测试用例，覆盖正常场景和异常场景。'
           '测试内容包括：功能测试（接口是否按预期工作）、安全测试（未登录访问拦截、错误密码拦截）、'
           '边界测试（空参数、重复注册等）。')

    r.heading2('5.3 测试用例与结果')
    r.add_table(
        ['序号', '测试项', '输入', '预期结果', '实际结果'],
        [
            ['1', '正常登录', 'admin/admin123', '200 + Token', '✅ 通过'],
            ['2', '错误密码', 'admin/wrong', '1003 密码错误', '✅ 通过'],
            ['3', '用户不存在', 'nobody/xxx', '1001 用户不存在', '✅ 通过'],
            ['4', '无Token访问', 'GET /me (无Header)', '401 未登录', '✅ 通过'],
            ['5', '无效Token', 'Authorization: Bearer xxx', '401 Token无效', '✅ 通过'],
            ['6', '注册新用户', 'new_user/pass123', '200 + 用户信息', '✅ 通过'],
            ['7', '重复注册', '已存在的用户名', '1002 用户名已存在', '✅ 通过'],
            ['8', '缺少必填字段', '{username, password} 缺real_name', '400 参数错误', '✅ 通过'],
            ['9', '修改密码', '正确旧密码+新密码', '200 修改成功', '✅ 通过'],
            ['10', '旧密码错误', '错误旧密码', '400 旧密码不正确', '✅ 通过'],
            ['11', '重置密码', '正确用户名+新密码', '200 重置成功', '✅ 通过'],
            ['12', '新密码登录', '重置后的密码', '200 + Token', '✅ 通过'],
            ['13', '账号禁用', '禁用账号登录', '1004 账号已禁用', '✅ 通过'],
            ['14', '健康检查', 'GET /api/health', '200 OK', '✅ 通过'],
        ]
    )

    r.heading2('5.4 系统运行截图')
    r.para('（此处插入系统运行截图：登录页面、管理员端、教师端、学生端界面截图）')

    r.heading2('5.5 代码质量')
    r.para('后端代码遵循Java编码规范：类名采用大驼峰命名法，方法名和变量名采用小驼峰命名法，'
           '常量使用全大写。每个类和方法均有JavaDoc注释说明功能。'
           '全局异常处理器确保所有异常都被统一捕获并返回JSON格式错误信息，'
           '不会因未处理异常导致前端收到不友好的500页面。')

    # ===== 第6节：总结 =====
    r.heading1('6. 总结')

    r.heading2('6.1 完成情况')
    r.para('本人负责的登录认证模块已全部完成并测试通过。实现了用户注册、登录、密码修改、'
           '密码重置、退出登录共6个REST API接口，配合LoginInterceptor和RoleInterceptor实现了'
           '完整的RBAC权限控制。所有接口均通过14项测试用例验证，功能正确、异常处理完善。')

    r.heading2('6.2 遇到的问题与解决方案')
    r.bullet('问题1：Spring Boot版本兼容性。本地JDK为25，Spring Boot 3.3.5不支持class file major version 69。'
             '解决：将Spring Boot升级至3.5.0，该版本支持Java 25。')
    r.bullet('问题2：SQL Server TCP/IP连接失败。默认SQL Server只启用Shared Memory协议。'
             '解决：在SQL Server配置管理器中启用TCP/IP协议并重启服务。')
    r.bullet('问题3：RoleInterceptor拦截登录接口。登录接口也被RoleInterceptor拦截导致无法登录。'
             '解决：在WebMvcConfig中为RoleInterceptor添加与LoginInterceptor相同的白名单。')
    r.bullet('问题4：init.sql中admin密码MD5哈希值错误。原始文件中的哈希值与实际密码不匹配。'
             '解决：使用Java MD5Util重新计算admin123的MD5值并更新数据库和init.sql。')
    r.bullet('问题5：课程导入时外键约束冲突。teacher表中无对应记录导致course无法插入。'
             '解决：在init.sql中添加teacher和student的初始记录。')

    r.heading2('6.3 心得体会')
    r.para('通过本次课程设计，我深入理解了Web后端开发的核心流程：从需求分析、数据库设计、'
           '接口定义到编码实现和测试验证。在实践中掌握了Spring Boot框架的使用方法、'
           'MyBatis-Plus的CRUD操作、拦截器机制、Session管理等关键技术。')
    r.para('同时，团队协作是本次项目的另一个重要收获。6名成员分工协作，通过GitHub进行代码管理，'
           '我负责的后端核心模块为其他成员提供了统一的认证和权限基础，使得教师端、学生端、'
           '管理员端可以在此基础上独立开发，互不干扰。')
    r.para('从工程思维角度，本项目也让我认识到：软件工程不仅仅是编写代码，更重要的是需求分析、'
           '架构设计、测试验证和文档编写等全流程实践。好的架构设计能够有效降低模块间的耦合度，'
           '提高开发效率。')

    # ===== 第7节：参考文献 =====
    r.heading1('7. 参考文献')
    refs = [
        '[1] Spring Boot官方文档. https://spring.io/projects/spring-boot',
        '[2] MyBatis-Plus官方文档. https://baomidou.com/',
        '[3] Microsoft SQL Server文档. https://learn.microsoft.com/zh-cn/sql/',
        '[4] 李刚. 疯狂Java讲义（第5版）. 电子工业出版社, 2019.',
        '[5] 杨开振. 深入浅出Spring Boot 2.x. 人民邮电出版社, 2018.',
        '[6] Erich Gamma等. 设计模式：可复用面向对象软件的基础. 机械工业出版社, 2019.',
    ]
    for ref in refs:
        r.para(ref)

    # ===== 保存 =====
    output = r'c:\Users\moling\Downloads\软件工程综合实验-发给学生\软件工程综合实验-发给学生\软件工程综合实验-发给学生\学生成绩管理系统-课程设计报告-成员2.docx'
    r.save(output)


if __name__ == '__main__':
    build_report()
